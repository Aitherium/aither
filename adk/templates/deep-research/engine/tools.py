"""Curated tool set for the Deep Research Agent.

These are the ONLY tools the agent may call (default-deny; the brain pack lists
them as a whitelist). Read & generate only — no shell, no file delete, no deploy.

Each tool is a closure over a ResearchSession so it can:
  - reach the knowledge graph (adk GraphMemory, SQLite, offline-safe),
  - feed the SavingsLedger (so the UI can show tokens used vs. saved),
  - share fetched-page and findings caches across the session (dedup).

`build_research_tools(session)` returns a list of plain functions; serve.py
registers them on the AitherAgent via agent._tools.register(fn).
"""

from __future__ import annotations

import json
import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urljoin, urlsplit

import httpx

from . import aithersearch
from .ledger import SavingsLedger

logger = logging.getLogger("deep_research.tools")


@dataclass
class ResearchSession:
    """Shared state for one research session: graph, ledger, caches, artifacts."""

    graph: Any                       # adk.graph_memory.GraphMemory
    ledger: SavingsLedger
    artifacts_dir: Path
    session_id: str = "session"
    _page_cache: dict[str, str] = field(default_factory=dict)
    _finding_keys: set[str] = field(default_factory=set)
    sources: list[dict] = field(default_factory=list)  # ordered, for citations

    def cite(self, title: str, url: str) -> int:
        """Register a source and return its 1-based citation number."""
        for i, s in enumerate(self.sources):
            if s["url"] == url:
                return i + 1
        self.sources.append({"title": title, "url": url})
        return len(self.sources)


def _strip_html(html: str) -> str:
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<[^>]+>", " ", html)
    html = re.sub(r"\s+", " ", html)
    return html.strip()


def _as_int(v, default: int) -> int:
    """Coerce an LLM-supplied arg to int. Models sometimes pass lists/strings/None
    (e.g. `angles=["overview","recent"]`); never crash on that."""
    if isinstance(v, bool):
        return default
    if isinstance(v, int):
        return v
    if isinstance(v, float):
        return int(v)
    if isinstance(v, (list, tuple, set)):
        return len(v) or default
    if isinstance(v, str):
        m = re.search(r"-?\d+", v)
        return int(m.group()) if m else default
    return default


_PRIVATE_HOST = re.compile(
    r"^(localhost|0\.0\.0\.0|127\.|10\.|192\.168\.|169\.254\.|::1|"
    r"172\.(1[6-9]|2\d|3[01])\.)|(\.internal|\.local)$", re.IGNORECASE)


def _safe_public_url(url: str) -> bool:
    """Allow only public http(s) URLs (block localhost/private/link-local hosts)."""
    try:
        p = urlsplit(url)
    except ValueError:
        return False
    host = (p.hostname or "").lower()
    return p.scheme in ("http", "https") and bool(host) and not _PRIVATE_HOST.search(host)


def _sitemap_locs(xml_text: str) -> list[str]:
    """Extract <loc> URLs from a sitemap or sitemap index (regex — robust enough)."""
    return re.findall(r"<loc>\s*([^<\s]+)\s*</loc>", xml_text)


def _kb_rel_path(url: str) -> str:
    """Map a URL to a safe relative .md path mirroring its tree."""
    path = urlsplit(url).path.strip("/") or "index"
    path = re.sub(r"\.(html?|md)$", "", path, flags=re.I)
    segs = [re.sub(r"[^A-Za-z0-9._-]+", "-", s).strip("-") or "_" for s in path.split("/")]
    return "/".join(segs) + ".md"


def build_research_tools(session: ResearchSession) -> list[Callable]:
    graph = session.graph
    ledger = session.ledger

    async def web_search(query: str, limit: int = 6) -> str:
        """Search the open web (DuckDuckGo, no API key). Returns titles, URLs, and snippets.

        query: what to search for
        limit: max results (default 6)
        """
        limit = _as_int(limit, 6)
        ledger.note_search()
        res = await aithersearch.search(str(query), limit=limit)
        # Register results as citable sources up front.
        for r in res.get("results", []):
            if r.get("url"):
                session.cite(r.get("title", ""), r["url"])
        return json.dumps(res)

    async def deep_research(topic: str, angles: int = 3) -> str:
        """Run a multi-angle web sweep on a topic and return de-duplicated results.

        Fires several complementary queries (overview, recent developments,
        comparisons/criticism) and merges the hits. Use for broad questions.

        topic: the subject to investigate
        angles: how many query angles to run (default 3, max 5)
        """
        angles = max(1, min(_as_int(angles, 3), 5))
        topic = str(topic)
        templates = [
            "{t}",
            "{t} latest developments 2026",
            "{t} comparison OR analysis",
            "{t} criticism OR limitations OR risks",
            "{t} data OR statistics OR report",
        ][:angles]
        merged: dict[str, dict] = {}
        for tmpl in templates:
            ledger.note_search()
            res = await aithersearch.search(tmpl.format(t=topic), limit=5)
            for r in res.get("results", []):
                u = r.get("url", "")
                if u and u not in merged:
                    merged[u] = r
                    session.cite(r.get("title", ""), u)
        return json.dumps({"topic": topic, "queries": len(templates),
                           "results": list(merged.values()), "count": len(merged)})

    async def fetch_url(url: str, max_chars: int = 6000) -> str:
        """Fetch a web page and return its readable text (HTML stripped).

        Reuses a per-session cache: re-fetching the same URL costs nothing and is
        credited as tokens saved.

        url: the page to read
        max_chars: max characters of text to return (default 6000)
        """
        max_chars = _as_int(max_chars, 6000)
        url = str(url)
        if url in session._page_cache:
            cached = session._page_cache[url]
            ledger.record_dedup(cached[:max_chars])  # avoided a re-fetch + re-read
            return json.dumps({"url": url, "text": cached[:max_chars], "cached": True})
        try:
            async with httpx.AsyncClient(timeout=15.0, follow_redirects=True,
                                         max_redirects=5) as client:
                resp = await client.get(url, headers={
                    "User-Agent": "Mozilla/5.0 (compatible; DeepResearchAgent/1.0)"})
                resp.raise_for_status()
                text = _strip_html(resp.text)
        except Exception as exc:  # noqa: BLE001
            return json.dumps({"url": url, "error": str(exc)})
        session._page_cache[url] = text
        ledger.note_page_read()
        return json.dumps({"url": url, "text": text[:max_chars], "cached": False,
                           "length": len(text)})

    async def mirror_website(base_url: str, max_pages: int = 25) -> str:
        """Mirror a whole docs/site into a local markdown knowledge base and cite it.

        Discovers pages under the URL (sitemap.xml + llms.txt, scoped to the same
        host + path prefix), downloads each as markdown (Mintlify ".md"-suffix when
        available, else HTML stripped to text), writes them under the artifacts dir,
        builds an INDEX.md, and registers every page as a citable source.

        Use this INSTEAD of many `fetch_url` calls when you need a whole docs site.

        base_url: root URL to mirror (e.g. "https://platform.claude.com/docs/en/api/")
        max_pages: safety cap on pages to mirror (default 25)
        """
        base_url = str(base_url).strip()
        max_pages = max(1, min(_as_int(max_pages, 25), 200))
        if not _safe_public_url(base_url):
            return json.dumps({"error": "URL blocked (must be a public http/https site)",
                               "base_url": base_url})

        parts = urlsplit(base_url)
        origin = f"{parts.scheme}://{parts.netloc}"
        prefix = parts.path.rstrip("/")
        domain = parts.netloc.replace(":", "_")
        out_dir = session.artifacts_dir / "mirrors" / domain
        ledger.note_search()

        headers = {"User-Agent": "Mozilla/5.0 (compatible; DeepResearchAgent/1.0)"}
        urls: list[str] = []
        try:
            async with httpx.AsyncClient(timeout=20.0, follow_redirects=True,
                                         max_redirects=5) as client:
                # 1) discover via sitemap + llms.txt
                for disc in (urljoin(origin, "/sitemap.xml"), urljoin(origin, "/llms.txt")):
                    try:
                        r = await client.get(disc, headers=headers)
                        if r.status_code != 200 or not r.text:
                            continue
                        if disc.endswith(".xml"):
                            urls += _sitemap_locs(r.text)
                        else:
                            urls += [u[:-3] if u.endswith(".md") else u
                                     for u in re.findall(r"https?://[^\s)\]]+", r.text)]
                    except httpx.HTTPError:
                        continue
                # scope to same host + path prefix; dedupe; cap
                seen: set[str] = set()
                scoped: list[str] = []
                for u in [base_url, *urls]:
                    pu = urlsplit(u)
                    norm = f"{pu.scheme}://{pu.netloc}{pu.path}"
                    if pu.netloc != parts.netloc or not pu.path.startswith(prefix or "/"):
                        continue
                    if re.search(r"\.(png|jpe?g|gif|svg|ico|css|js|pdf|zip|woff2?)$", norm, re.I):
                        continue
                    if norm in seen or not _safe_public_url(norm):
                        continue
                    seen.add(norm)
                    scoped.append(norm)
                    if len(scoped) >= max_pages:
                        break

                # 2) download each as markdown (.md-suffix first, else HTML->text)
                pages: list[dict] = []
                for u in scoped:
                    md, title, src = "", u.rsplit("/", 1)[-1], "html"
                    try:
                        if not u.endswith(".md"):
                            rm = await client.get(u + ".md", headers=headers)
                            body = rm.text if rm.status_code == 200 else ""
                            if body and not body.lstrip()[:64].lower().startswith(
                                    ("<!doctype", "<html", "<?xml")):
                                md, src = body.strip(), "markdown"
                        if not md:
                            rp = await client.get(u, headers=headers)
                            if rp.status_code == 200:
                                md = _strip_html(rp.text)
                                tm = re.search(r"<title[^>]*>([^<]+)</title>", rp.text, re.I)
                                if tm:
                                    title = tm.group(1).strip()
                    except httpx.HTTPError:
                        continue
                    if not md:
                        continue
                    for line in md.splitlines():
                        if line.startswith("# "):
                            title = line[2:].strip()
                            break
                    dest = out_dir / _kb_rel_path(u)
                    dest.parent.mkdir(parents=True, exist_ok=True)
                    dest.write_text(f"---\nurl: {u}\ntitle: {title}\n---\n\n{md}\n",
                                    encoding="utf-8")
                    n = session.cite(title, u)
                    ledger.note_page_read()
                    pages.append({"url": u, "title": title, "file": _kb_rel_path(u),
                                  "source": src, "cite": n})
        except Exception as exc:  # noqa: BLE001
            return json.dumps({"error": str(exc), "base_url": base_url})

        if not pages:
            return json.dumps({"error": "no pages mirrored", "base_url": base_url,
                               "discovered": len(scoped)})
        # 3) INDEX.md
        idx = [f"# Mirror: {domain}", "", f"Source: {base_url}",
               f"Pages: {len(pages)}", "", "## Pages", ""]
        idx += [f"- [{p['title']}]({p['file']}) — [{p['cite']}]" for p in pages]
        (out_dir / "INDEX.md").write_text("\n".join(idx) + "\n", encoding="utf-8")

        return json.dumps({
            "success": True, "base_url": base_url,
            "pages_discovered": len(scoped), "pages_mirrored": len(pages),
            "output_dir": str(out_dir), "index_file": str(out_dir / "INDEX.md"),
            "sources_registered": len(pages),
            "note": "All pages are now citable with [n]; see INDEX.md for the list.",
        })

    async def save_finding(claim: str, source_url: str = "", topic: str = "") -> str:
        """Store a verified fact + its source in the knowledge graph for reuse.

        Call this for every solid fact you find so you never re-research it.

        claim: the fact, in one sentence
        source_url: the URL it came from
        topic: optional grouping label
        """
        key = re.sub(r"\s+", " ", claim.strip().lower())[:160]
        if key in session._finding_keys:
            ledger.record_dedup(claim)  # already known — no re-store/re-embed
            return json.dumps({"stored": False, "duplicate": True})
        session._finding_keys.add(key)
        tags = [t for t in ["finding", topic] if t]
        try:
            await graph.add_node(
                label=(topic or claim[:60]),
                node_type="fact",
                content=f"{claim} [source: {source_url}]" if source_url else claim,
                tags=tags,
                source_session=session.session_id,
            )
        except Exception as exc:  # noqa: BLE001
            logger.debug("graph.add_node failed: %s", exc)
        ledger.note_finding()
        cite_n = session.cite("", source_url) if source_url else 0
        return json.dumps({"stored": True, "citation": cite_n})

    async def recall(query: str, limit: int = 5) -> str:
        """Recall what you already learned this session from memory + the graph.

        ALWAYS call this before searching — if you already know it, reuse it and
        cite the cached finding instead of spending another search. Reused tokens
        are counted as 'saved by memory'.

        query: what you're trying to remember
        limit: max results (default 5)
        """
        limit = _as_int(limit, 5)
        try:
            nodes = await graph.search(str(query), limit=limit)
        except Exception as exc:  # noqa: BLE001
            return json.dumps({"results": [], "error": str(exc)})
        results = []
        reused_text = ""
        for n in nodes:
            item = {"label": n.label, "content": n.content, "type": n.node_type}
            results.append(item)
            reused_text += " " + (n.content or "")
        if results:
            ledger.record_recall(reused_text)  # answered from memory, not a fresh fetch
        return json.dumps({"results": results, "count": len(results),
                           "reused_from_memory": bool(results)})

    async def knowledge_graph() -> str:
        """Show what the agent has learned this session (graph stats + counts)."""
        try:
            stats = await graph.get_stats()
        except Exception as exc:  # noqa: BLE001
            stats = {"error": str(exc)}
        return json.dumps({"graph": stats, "sources_cited": len(session.sources)})

    # ── Report generation (read/generate only; writes to the artifacts dir) ──

    def _sources_block() -> str:
        if not session.sources:
            return ""
        lines = ["", "## Sources", ""]
        for i, s in enumerate(session.sources, 1):
            title = s.get("title") or s.get("url")
            lines.append(f"{i}. {title} — {s.get('url')}")
        return "\n".join(lines)

    def _save(name: str, data: bytes) -> Path:
        session.artifacts_dir.mkdir(parents=True, exist_ok=True)
        ts = time.strftime("%Y%m%d-%H%M%S")
        safe = re.sub(r"[^a-zA-Z0-9_.-]+", "-", name).strip("-") or "report"
        path = session.artifacts_dir / f"{ts}-{safe}"
        path.write_bytes(data)
        return path

    def generate_markdown(title: str, content: str) -> str:
        """Write a Markdown research report (auto-appends a Sources section).

        title: report title
        content: report body in Markdown (use [1], [2] inline citations)
        """
        md = f"# {title}\n\n{content}\n{_sources_block()}\n"
        path = _save(f"{title}.md", md.encode("utf-8"))
        return json.dumps({"saved": True, "format": "markdown",
                           "path": str(path), "file": path.name})

    def generate_pdf(title: str, content: str) -> str:
        """Write a PDF research report (reportlab). Auto-appends Sources.

        title: report title
        content: report body in Markdown-ish text (#, ##, - bullets, [1] citations)
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            return json.dumps({"saved": False, "error": "reportlab not installed"})

        import io
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter, title=title,
                                topMargin=0.8 * inch, bottomMargin=0.8 * inch)
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("H1c", parent=styles["Heading1"], spaceAfter=12)
        story = [Paragraph(title, h1), Spacer(1, 8)]
        body = content + "\n" + _sources_block()
        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"&bull;&nbsp;{_esc(line[2:])}", styles["BodyText"]))
            else:
                story.append(Paragraph(_esc(line), styles["BodyText"]))
        doc.build(story)
        path = _save(f"{title}.pdf", buf.getvalue())
        return json.dumps({"saved": True, "format": "pdf",
                           "path": str(path), "file": path.name})

    def generate_docx(title: str, content: str) -> str:
        """Write a Word (.docx) research report (python-docx). Auto-appends Sources.

        title: report title
        content: report body (#, ##, - bullets, [1] citations)
        """
        try:
            from docx import Document
        except ImportError:
            return json.dumps({"saved": False, "error": "python-docx not installed"})
        import io
        document = Document()
        document.add_heading(title, level=0)
        body = content + "\n" + _sources_block()
        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line:
                continue
            if line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)
        buf = io.BytesIO()
        document.save(buf)
        path = _save(f"{title}.docx", buf.getvalue())
        return json.dumps({"saved": True, "format": "docx",
                           "path": str(path), "file": path.name})

    return [
        web_search, deep_research, fetch_url, mirror_website, save_finding, recall,
        knowledge_graph, generate_markdown, generate_pdf, generate_docx,
    ]


def _esc(text: str) -> str:
    """Escape the few chars reportlab's mini-markup cares about."""
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
